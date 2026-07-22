"""
Build Word (.docx) copies of the technical docs, with Mermaid diagrams rendered
to embedded images.

    python scripts/build_word_docs.py

Renders each mermaid code block to a PNG via @mermaid-js/mermaid-cli (through
npx, which downloads a headless Chromium on first run), then assembles a
formatted .docx with python-docx. No Microsoft Word installation is required.

Self-contained on purpose: the Markdown-to-docx converter lives in this file so
the whole generator is one committed script. It is not a general CommonMark
parser — it handles exactly the constructs the AirWatch docs use (ATX headings,
pipe tables, fenced code, images, lists, blockquotes, inline bold/italic/code).
"""
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
DOCS = REPO / "docs"
OUT = DOCS / "word"

# Rendered from docs/<stem>.md into docs/word/<stem>.docx.
TARGETS = ["TECHNICAL_DOCUMENT", "ARCHITECTURE"]

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"


# ─── Markdown → docx ──────────────────────────────────────────────────────────

def _shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def add_inline(paragraph, text):
    token = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?\*|`[^`]+`)")
    pos = 0
    for m in token.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        piece = m.group(0)
        if piece.startswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`"):
            r = paragraph.add_run(piece[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
        else:
            paragraph.add_run(piece[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    _shade(p._p.get_or_add_pPr(), CODE_BG)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = CODE_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_table(doc, rows):
    header, body = rows[0], rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], cell_text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].paragraphs[0].text = ""
            add_inline(cells[i].paragraphs[0], r[i] if i < len(r) else "")
    doc.add_paragraph()


def _row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _rule(doc):
    p = doc.add_paragraph()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def convert(md_path, out_path, img_base):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = BODY_FONT
    doc.styles["Normal"].font.size = Pt(10.5)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, buf)
            continue

        m = re.match(r"!\[.*?\]\((.+?)\)", line.strip())
        if m:
            p = Path(m.group(1))
            if not p.is_absolute():
                p = Path(img_base) / p
            if p.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    para.add_run().add_picture(str(p), width=Inches(6.2))
                except Exception:
                    para.add_run(f"[diagram: {m.group(1)}]")
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            heading.text = ""
            content = re.sub(r"[#*`]", "", m.group(2)) if level == 1 else m.group(2)
            add_inline(heading, content)
            for run in heading.runs:
                run.font.color.rgb = ACCENT
            i += 1
            continue

        if re.match(r"^\s*---+\s*$", line):
            _rule(doc)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if line.strip().startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            pbdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "F97316")
            pbdr.append(left)
            p._p.get_or_add_pPr().append(pbdr)
            add_inline(p, " ".join(x.strip() for x in buf))
            for run in p.runs:
                run.italic = True
            continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            add_inline(doc.add_paragraph(style="List Number"), m.group(3))
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            if len(m.group(1)) >= 2:
                p.paragraph_format.left_indent = Inches(0.6)
            add_inline(p, m.group(2))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        buf = [line]
        i += 1
        stop = re.compile(r"^(#{1,6}\s|\s*[-*]\s|\s*\d+\.\s|>|\|\s*|```|!\[)")
        while i < n and lines[i].strip() and not stop.match(lines[i]) and not re.match(r"^\s*---+\s*$", lines[i]):
            buf.append(lines[i])
            i += 1
        add_inline(doc.add_paragraph(), " ".join(x.strip() for x in buf))

    doc.save(out_path)
    print(f"wrote {out_path}")


# ─── Render + orchestrate ─────────────────────────────────────────────────────

def render_and_build(stem, workdir):
    src = DOCS / f"{stem}.md"
    if not src.exists():
        print(f"skip {stem}: {src} missing")
        return False

    (workdir / "mmdc.json").write_text(
        json.dumps({"theme": "default", "flowchart": {"useMaxWidth": True}}), encoding="utf-8"
    )
    (workdir / "puppeteer.json").write_text(json.dumps({"args": ["--no-sandbox"]}), encoding="utf-8")

    rendered_md = workdir / f"{stem}.md"
    print(f"=== {stem}: rendering diagrams (first run downloads Chromium) ===")
    cmd = [
        "npx", "-y", "@mermaid-js/mermaid-cli@11",
        "-i", str(src), "-o", str(rendered_md),
        "-b", "white",
        "-c", str(workdir / "mmdc.json"),
        "-p", str(workdir / "puppeteer.json"),
        "-e", "png",
    ]
    # shell=True so Windows resolves the npx.cmd shim.
    subprocess.run(" ".join(f'"{c}"' if " " in str(c) else str(c) for c in cmd), shell=True, check=True)

    OUT.mkdir(parents=True, exist_ok=True)
    out_docx = OUT / f"{stem}.docx"
    print(f"=== {stem}: building {out_docx.name} ===")
    convert(str(rendered_md), str(out_docx), str(workdir))
    return True


def main():
    built = []
    with tempfile.TemporaryDirectory() as tmp:
        for stem in TARGETS:
            try:
                if render_and_build(stem, Path(tmp)):
                    built.append(stem)
            except subprocess.CalledProcessError as exc:
                print(f"FAILED to render {stem}: {exc}", file=sys.stderr)
    print(f"\nBuilt {len(built)}: {', '.join(built)} -> {OUT}")
    return 0 if built else 1


if __name__ == "__main__":
    raise SystemExit(main())
